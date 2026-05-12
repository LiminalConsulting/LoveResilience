"""Build a Word template for Tanja's Love Resilience booklet.

Spec from the printer:
  - Page size: 85 x 125 mm
  - Margins: 15 mm Bund (inside/spine), 10 mm outside, top, bottom
  - Text area (Satzspiegel): 60 x 105 mm
  - Mirror margins for facing pages
  - Target: ~200 pages, font size chosen to fit
"""

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT_PATH = "/Users/davidrug/RealDealVault/LiminalConsulting/LoveResilience-Buechlein/LoveResilience-Buechlein-Vorlage.docx"

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Mm(85)
section.page_height = Mm(125)
section.top_margin = Mm(10)
section.bottom_margin = Mm(10)
# In Word, when mirror margins are on:
#   left_margin  = inside (Bund)
#   right_margin = outside
# (Word will mirror these automatically on even pages.)
section.left_margin = Mm(15)
section.right_margin = Mm(10)
section.gutter = Mm(0)
section.header_distance = Mm(5)
section.footer_distance = Mm(5)

# Enable mirror margins + different-odd-even-pages on the section
sectPr = section._sectPr

# mirror margins
mirror = OxmlElement('w:mirrorMargins')
sectPr.insert(0, mirror)

# Document-level settings: language + autoHyphenation + evenAndOddHeaders
settings = doc.settings.element  # <w:settings>

def add_or_replace(parent, tag, attrs=None):
    el = parent.find(qn(tag))
    if el is not None:
        parent.remove(el)
    el = OxmlElement(tag)
    if attrs:
        for k, v in attrs.items():
            el.set(qn(k), v)
    parent.append(el)
    return el

# Auto-hyphenation on
add_or_replace(settings, 'w:autoHyphenation')
# Hyphenate caps off (cleaner for body text)
add_or_replace(settings, 'w:doNotHyphenateCaps')
# Consecutive hyphens limit (3 in a row max — typographic best practice)
add_or_replace(settings, 'w:consecutiveHyphenLimit', {'w:val': '3'})
# Hyphenation zone (smaller = tighter hyphenation, default 360 twips)
add_or_replace(settings, 'w:hyphenationZone', {'w:val': '283'})  # ~5mm
# Different odd/even headers (lets us mirror running heads later if wanted)
add_or_replace(settings, 'w:evenAndOddHeaders')

# --- Default font + language for the Normal style ---
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Cambria'
normal.font.size = Pt(10)
# Set East-Asian/complex-script font names so Word doesn't fall back
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.insert(0, rfonts)
rfonts.set(qn('w:ascii'), 'Cambria')
rfonts.set(qn('w:hAnsi'), 'Cambria')
rfonts.set(qn('w:cs'), 'Cambria')

# Language = German (de-DE)
lang = rpr.find(qn('w:lang'))
if lang is None:
    lang = OxmlElement('w:lang')
    rpr.append(lang)
lang.set(qn('w:val'), 'de-DE')
lang.set(qn('w:eastAsia'), 'de-DE')
lang.set(qn('w:bidi'), 'de-DE')

# --- Paragraph defaults for Normal: justified, 1.15 line spacing, small first-line indent ---
ppr = normal.element.get_or_add_pPr()

# justification
jc = ppr.find(qn('w:jc'))
if jc is None:
    jc = OxmlElement('w:jc')
    ppr.append(jc)
jc.set(qn('w:val'), 'both')

# line spacing 1.15 (276 twentieths of a point per line is standard for 1.15x of single)
spacing = ppr.find(qn('w:spacing'))
if spacing is None:
    spacing = OxmlElement('w:spacing')
    ppr.append(spacing)
spacing.set(qn('w:line'), '276')
spacing.set(qn('w:lineRule'), 'auto')
spacing.set(qn('w:after'), '0')
spacing.set(qn('w:before'), '0')

# First-line indent: 3mm (gentle, book-like). Comment out for no indent.
ind = ppr.find(qn('w:ind'))
if ind is None:
    ind = OxmlElement('w:ind')
    ppr.append(ind)
ind.set(qn('w:firstLine'), str(int(3 * 56.6929)))  # 3mm in twips (1mm ≈ 56.69 twips)

# --- Add a starter paragraph so Tanja sees the layout immediately ---
p = doc.add_paragraph(
    "Hier deinen Text einfügen oder eintippen. Die Seite ist bereits auf "
    "85 × 125 mm eingestellt; die Ränder (15 mm innen am Bund, 10 mm außen, "
    "oben und unten) sind automatisch gespiegelt für linke und rechte Seiten. "
    "Du kannst diesen Hinweistext einfach löschen und mit deinem Inhalt "
    "weiterschreiben."
)

# A second paragraph so the justify behaviour is visible right away
p2 = doc.add_paragraph(
    "Der Blocksatz, die Silbentrennung (deutsch) und die Schriftgröße (10 pt) "
    "sind bereits eingestellt. Wenn der Text zu viele Seiten wird, kannst du "
    "über Format → Schriftart die Größe anpassen — zum Beispiel auf 9 pt."
)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
