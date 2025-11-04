#!/usr/bin/env python3
"""
Generate a client document (DOCX format compatible with Apple Pages)
that consolidates all card markdown files with embedded images.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_file(filepath):
    """Parse a markdown file and extract frontmatter and sections."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract frontmatter
    frontmatter_match = re.match(r'^---\n(.*?)\n---\n(.*)', content, re.DOTALL)
    if not frontmatter_match:
        raise ValueError(f"Invalid markdown format in {filepath}")

    frontmatter_text = frontmatter_match.group(1)
    body = frontmatter_match.group(2)

    # Parse frontmatter
    frontmatter = {}
    for line in frontmatter_text.split('\n'):
        if ':' in line:
            key, value = line.split(':', 1)
            frontmatter[key.strip()] = value.strip().strip('"')

    # Extract sections
    sections = {}
    current_section = None
    current_content = []

    for line in body.split('\n'):
        if line.startswith('# '):
            if current_section:
                sections[current_section] = '\n'.join(current_content).strip()
            current_section = line[2:].strip()
            current_content = []
        else:
            current_content.append(line)

    if current_section:
        sections[current_section] = '\n'.join(current_content).strip()

    return frontmatter, sections

def add_card_to_document(doc, frontmatter, sections, image_path=None):
    """Add a card's content to the document with formatting."""

    # Card title (name)
    title = doc.add_heading(frontmatter.get('name', 'Untitled'), level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add card image if it exists
    if image_path and os.path.exists(image_path):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        try:
            # Add image with width of 3 inches (adjust as needed)
            run.add_picture(str(image_path), width=Inches(3))
        except Exception as e:
            print(f"Warning: Could not add image {image_path}: {e}")

    # Add each section
    section_order = ['Meaning', 'Question', 'Action', 'Quote']
    for section_name in section_order:
        if section_name in sections:
            # Section heading
            heading = doc.add_heading(section_name, level=2)

            # Section content
            content = sections[section_name]

            # Handle quotes specially (they start with >)
            if section_name == 'Quote':
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('>'):
                        # Remove the > and process
                        quote_text = line[1:].strip()
                        p = doc.add_paragraph(quote_text)
                        p.paragraph_format.left_indent = Inches(0.5)
                        # Italicize quote
                        for run in p.runs:
                            run.italic = True
                    elif line:  # Non-empty line
                        p = doc.add_paragraph(line)
                        p.paragraph_format.left_indent = Inches(0.5)
            else:
                # Regular paragraph
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())

    # Add page break after each card
    doc.add_page_break()

def generate_document(cards_dir, images_dir, output_path):
    """Generate the complete client document."""

    # Create document
    doc = Document()

    # Set document title
    title = doc.add_heading('Love Resilience Card Deck', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Spacing

    # Get all markdown files sorted by number
    cards_dir = Path(cards_dir)
    images_dir = Path(images_dir)

    md_files = sorted(cards_dir.glob('*.md'),
                     key=lambda x: int(re.search(r'(\d+)_', x.name).group(1)))

    print(f"Found {len(md_files)} card files")

    # Process each card
    for md_file in md_files:
        print(f"Processing: {md_file.name}")

        try:
            frontmatter, sections = parse_markdown_file(md_file)

            # Find corresponding image
            image_filename = frontmatter.get('image')
            image_path = None
            if image_filename:
                image_path = images_dir / image_filename
                if not image_path.exists():
                    print(f"  Warning: Image not found: {image_path}")
                    image_path = None

            # Add card to document
            add_card_to_document(doc, frontmatter, sections, image_path)

        except Exception as e:
            print(f"  Error processing {md_file.name}: {e}")
            continue

    # Save document
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024 / 1024:.2f} MB")

if __name__ == '__main__':
    # Paths
    CARDS_DIR = 'card-content'
    IMAGES_DIR = 'public/CardSet'
    OUTPUT_PATH = 'LoveResilience_CardDeck_Complete.docx'

    # Generate the document
    generate_document(CARDS_DIR, IMAGES_DIR, OUTPUT_PATH)
